from openai import OpenAI
import openpyxl


client = OpenAI(
    api_key='API KEY FOR OPEN AI'
)

wb = openpyxl.load_workbook("QUIZ FILE.xlsx")
ws = wb.active
responses = []
for row in ws.iter_rows(values_only=True):
    
    question = row[0]

    test_q = question

    completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": "You are a law student at an law school in the united states"},
        {"role": "user", "content": test_q}
    ]
    )
    
    response = completion.choices[0].message.content

    responses.append(response)




import docx

doc = docx.Document()

p = doc.add_paragraph()

p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = 0

run = p.add_run(f'QUIZ FILE NAME Results:')
run.font.name = 'Calibri'
run.font.size = docx.shared.Pt(20)
doc.add_paragraph()

i = 1
for val in responses:
    p = doc.add_paragraph()
    run = p.add_run(f'{i})')


    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(20)
    p = doc.add_paragraph()
    run = p.add_run(val)
    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(11)
    doc.add_paragraph()
    i += 1

doc.save('QUIZ FILE NAME Results.docx')