from openai import OpenAI
import openpyxl


client = OpenAI(
    api_key='API KEY FOR OPEN AI'
)

wb = openpyxl.load_workbook("QUIZ FILE.xlsx")
ws = wb.active
responses = []
first_answers = []
second_answers = []
third_answers = []
correct_answers = []
for row in ws.iter_rows(values_only=True):
    
    question = row[0]
    ans_A = row[1]
    ans_B = row[2]
    ans_C = row[3]
    ans_D = row[4]
    corr_ans = row[5]
    correct_answers.append(corr_ans)
    test_q = question + '\nA) ' + ans_A + '\nB) ' + ans_B + '\nC) ' + ans_C + '\nD) ' + ans_D + '\n\nSurround the letter of each answer with square brackets and no spaces. Include an explanation of why the answer is correct. Begin the response with all three answers.'

    completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": "You are a law student at an law school in the united states"},
        {"role": "user", "content": test_q}
    ]
    )
    
    response = completion.choices[0].message.content
    first_ans_loc = response.find('[') + 1
    post_first = response[first_ans_loc + 1:]
    second_ans_loc = post_first.find('[') + 1
    post_second = post_first[second_ans_loc + 1:]
    third_ans_loc = post_second.find('[') + 1

    #GPT messes up the formatting sometimes, this is to catch it
    first_ans = response[first_ans_loc]
    if first_ans == ' ':
        first_ans = response[first_ans_loc + 1]
    second_ans = post_first[second_ans_loc]
    if second_ans == ' ':
        second_ans = post_first[second_ans_loc + 1]
    third_ans = post_second[third_ans_loc]
    if third_ans == ' ':
        third_ans = post_second[third_ans_loc + 1]

    first_answers.append(first_ans)
    second_answers.append(second_ans)
    third_answers.append(third_ans)
    

    responses.append(response)


print(correct_answers)
print(first_answers)
print(second_answers)
print(third_answers)

corr = 0
for i in range(len(correct_answers)):
    if correct_answers[i] == first_answers[i]:
        corr += 1

grade = (corr / len(correct_answers)) * 100
grade = str(round(grade, 2))




import docx

doc = docx.Document()

p = doc.add_paragraph()

p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = 0

run = p.add_run(f'QUIZ FILE NAME: {grade}%')
run.font.name = 'Calibri'
run.font.size = docx.shared.Pt(20)
doc.add_paragraph()

i = 1
for val in responses:
    p = doc.add_paragraph()

    if first_answers[i-1] == correct_answers[i-1]:
        run = p.add_run(f'{i}) Correct')
    elif second_answers[i-1] == correct_answers[i-1]:
        run = p.add_run(f'{i}) Incorrect, Second Answer Correct')
    elif third_answers[i-1] == correct_answers[i-1]:
        run = p.add_run(f'{i}) Incorrect, Third Answer Correct')
    else:
        run = p.add_run(f'{i}) Incorrect, Correct Answer not in top 3')


    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(20)
    p = doc.add_paragraph()
    run = p.add_run(val)
    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(11)
    doc.add_paragraph()
    i += 1

doc.save('QUIZ FILE NAME Responses.docx')